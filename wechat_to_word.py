"""
微信公众号文章保存为 Word 文档工具
功能：输入公众号文章链接，自动解析并保存为格式化的 Word 文档
优化版：增强提取稳定性、Word排版质量、UI体验、异常处理和知识库索引
"""

import sys
import os
import re
import io
import json
import subprocess
import traceback
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog,
    QProgressBar, QGroupBox, QMessageBox, QListWidget, QListWidgetItem,
    QSplitter, QShortcut,
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor, QPalette, QIcon, QKeySequence


# ==================== 网络与请求 ====================

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/125.0.0.0 Safari/537.36 MicroMessenger/7.0.20"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    "Referer": "https://mp.weixin.qq.com/",
}

RETRY_DELAYS = [1, 2, 3]


def fetch_with_retry(url, headers=None, timeout=30, retries=3):
    """带重试的 HTTP GET，返回 Response 对象"""
    headers = headers or HEADERS
    last_err = None
    for i in range(retries):
        try:
            resp = requests.get(url, headers=headers, timeout=timeout)
            resp.raise_for_status()
            return resp
        except requests.exceptions.Timeout:
            last_err = f"请求超时（第 {i+1} 次）"
        except requests.exceptions.ConnectionError:
            last_err = f"连接失败（第 {i+1} 次），请检查网络"
        except requests.exceptions.HTTPError as e:
            code = e.response.status_code if e.response else "?"
            if code == 404:
                raise ValueError("文章不存在（404），可能已被删除或链接无效。")
            if code in (403, 451):
                raise ValueError(f"访问被拒绝（{code}），文章可能需要登录或已被限制访问。")
            last_err = f"HTTP 错误 {code}（第 {i+1} 次）"
        except requests.exceptions.RequestException as e:
            last_err = f"请求异常（第 {i+1} 次）: {e}"
        if i < retries - 1:
            import time
            time.sleep(RETRY_DELAYS[i])
    raise ConnectionError(f"网络请求失败（已重试 {retries} 次）: {last_err}")


def fetch_article(url):
    """获取并解析公众号文章页面"""
    resp = fetch_with_retry(url)
    resp.encoding = "utf-8"
    return BeautifulSoup(resp.text, "lxml")


# ==================== 文章信息提取 ====================

def extract_title(soup):
    """提取文章标题，多级 fallback"""
    og = soup.find("meta", property="og:title")
    if og and og.get("content"):
        return og["content"].strip()

    h1 = soup.find("h1", id="activity-name")
    if h1:
        return h1.get_text(strip=True)

    h1 = soup.find("h1")
    if h1:
        return h1.get_text(strip=True)

    if soup.title:
        return soup.title.get_text(strip=True)

    return "未命名文章"


def extract_author(soup):
    """提取作者/公众号名称"""
    for sel in [
        ("strong", {"id": "profileBt"}),
        ("a", {"id": "js_name"}),
        ("meta", {"name": "author"}),
    ]:
        tag, attrs = sel
        el = soup.find(tag, attrs)
        if el:
            return (el.get("content") or el.get_text()).strip()

    # 从 profile_nickname 区域查找
    nick = soup.find("span", class_="profile_nickname")
    if nick:
        return nick.get_text(strip=True)

    return ""


def extract_date(soup):
    """提取发布日期，兼容时间戳和文本格式"""
    # 从 script 中提取时间戳
    for s in soup.find_all("script"):
        text = s.string or ""
        for pattern in [r'var\s+ct\s*=\s*["\']?(\d+)["\']?',
                        r'createTime\s*=\s*["\']?(\d+)["\']?',
                        r'publish_time\s*=\s*["\']?(\d+)["\']?']:
            m = re.search(pattern, text)
            if m:
                ts_str = m.group(1)
                ts = int(ts_str[:10]) if len(ts_str) >= 10 else int(ts_str)
                try:
                    dt = datetime.fromtimestamp(ts)
                    if 2010 <= dt.year <= 2099:
                        return dt.strftime("%Y-%m-%d")
                except (ValueError, OSError):
                    pass

    # em#publish_time
    em = soup.find("em", id="publish_time")
    if em:
        txt = em.get_text(strip=True)
        if txt:
            return txt

    # og:article:published_time
    og = soup.find("meta", property="og:article:published_time")
    if og and og.get("content"):
        return og["content"][:10]

    return ""


def extract_content_area(soup):
    """提取正文内容区域"""
    area = soup.find("div", id="js_content")
    if area:
        return area

    area = soup.find("div", class_="rich_media_content")
    if area:
        return area

    return None


def clean_content_area(area):
    """清理正文区域：移除隐藏元素、脚本、广告等无关内容"""
    # 移除隐藏元素
    for el in area.find_all(style=re.compile(r'display\s*:\s*none|visibility\s*:\s*hidden')):
        el.decompose()

    # 移除 script / style / iframe
    for tag_name in ("script", "style", "iframe", "noscript"):
        for el in area.find_all(tag_name):
            el.decompose()

    # 移除微信特有广告/推荐区域
    for sel in [
        {"id": "js_pc_qr_code"},
        {"id": "content_bottom_area"},
        {"class_": "reward_area"},
        {"class_": "rich_media_tool"},
        {"id": "js_toobar3"},
        {"class_": "qr_code_pc"},
        {"id": "js_profile_qrcode"},
    ]:
        for el in area.find_all(**sel):
            el.decompose()

    # 移除点赞/在看/阅读原文等尾部区域
    for el in area.find_all("div", id=re.compile(r"(js_toobar|js_share)")):
        el.decompose()

    return area


# ==================== 图片下载 ====================

def normalize_img_url(url):
    """补全图片 URL 协议"""
    if not url or url.startswith("data:"):
        return url
    if url.startswith("//"):
        return "https:" + url
    if not url.startswith("http"):
        return "https://" + url
    return url


def detect_image_format(data):
    """根据文件头检测图片格式"""
    if data[:3] == b'\xff\xd8\xff':
        return "jpeg"
    if data[:8] == b'\x89PNG\r\n\x1a\n':
        return "png"
    if data[:4] == b'GIF8':
        return "gif"
    if data[:4] == b'RIFF' and data[8:12] == b'WEBP':
        return "webp"
    return "png"


def download_image(url, downloaded=None, cache=None):
    """
    下载图片，返回 (bytes, format, actual_url)
    downloaded: set of already-downloaded URLs for dedup
    cache: dict of {url: (data, fmt)} for batch download cache
    """
    if downloaded is None:
        downloaded = set()

    url = normalize_img_url(url)
    if not url or url.startswith("data:"):
        return None, None, url

    # 优先从缓存取
    if cache and url in cache:
        data, fmt = cache[url]
        if data is not None:
            downloaded.add(url)
        return data, fmt, url

    # 去重
    if url in downloaded:
        return None, None, url

    try:
        resp = fetch_with_retry(url, timeout=15, retries=2)
        data = resp.content
        if len(data) < 100:
            return None, None, url

        fmt = detect_image_format(data)
        downloaded.add(url)
        return data, fmt, url
    except Exception:
        return None, None, url


def collect_image_urls(element):
    """递归收集元素中所有图片 URL"""
    urls = set()
    if isinstance(element, NavigableString):
        return urls
    for img in element.find_all("img"):
        src = get_img_src(img)
        if src:
            urls.add(normalize_img_url(src))
    return urls


def batch_download_images(urls, max_workers=8):
    """并发批量下载图片，返回 {url: (data, fmt)} 缓存"""
    cache = {}
    if not urls:
        return cache

    valid_urls = [u for u in urls if u and not u.startswith("data:")]

    def _download(url):
        try:
            resp = fetch_with_retry(url, timeout=15, retries=2)
            data = resp.content
            if len(data) < 100:
                return url, (None, None)
            fmt = detect_image_format(data)
            return url, (data, fmt)
        except Exception:
            return url, (None, None)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_download, u): u for u in valid_urls}
        for future in as_completed(futures):
            url, result = future.result()
            cache[url] = result

    return cache


def get_img_src(el):
    """从 img 元素获取图片 URL，兼容懒加载"""
    return (el.get("data-src")
            or el.get("data-original")
            or el.get("src")
            or "")


# ==================== Word 样式 ====================

def setup_doc_styles(doc):
    """设置文档基础样式"""
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.5
    # 设置中文字体
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    heading_sizes = {0: 22, 1: 20, 2: 16, 3: 14, 4: 13}
    for level, size in heading_sizes.items():
        sname = f"Heading {level}" if level > 0 else "Title"
        if sname in doc.styles:
            h = doc.styles[sname]
            h.font.name = "微软雅黑"
            h.font.size = Pt(size)
            h.font.color.rgb = RGBColor(30, 30, 30)
            h.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)


def add_hyperlink(paragraph, text, url):
    """在段落中添加可点击的超链接"""
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1155CC")
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(17, 85, 204)
        run.underline = True


def add_image_to_doc(doc, img_data, fmt, downloaded=None, img_counter=None):
    """将图片添加到文档，居中显示"""
    if img_data is None:
        return
    try:
        # webp 转换为 png
        if fmt == "webp":
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(img_data))
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                img_data = buf.getvalue()
                fmt = "png"
            except ImportError:
                doc.add_paragraph("[webp 图片需要安装 Pillow 库才能转换，请运行: pip install Pillow]")
                return
            except Exception as e:
                doc.add_paragraph(f"[webp 图片转换失败: {e}]")
                return

        ext = "jpg" if fmt == "jpeg" else fmt
        tmp = io.BytesIO(img_data)
        tmp.name = f"img.{ext}"

        doc.add_picture(tmp, width=Inches(5.8))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 图片编号
        if img_counter is not None:
            img_counter[0] += 1
            cap = doc.add_paragraph(f"图 {img_counter[0]}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
    except Exception as e:
        doc.add_paragraph(f"[图片加载失败: {e}]")


# ==================== 文本样式处理 ====================

def get_text_style(el):
    """分析元素的文本样式"""
    style = {}
    inline = el.get("style", "")
    if "font-weight" in inline or "bold" in inline:
        style["bold"] = True
    if "font-style" in inline or "italic" in inline:
        style["italic"] = True
    if "text-decoration" in inline and "underline" in inline:
        style["underline"] = True
    if "color" in inline:
        m = re.search(r"color\s*:\s*([^;]+)", inline)
        if m:
            style["color"] = m.group(1).strip()
    if "font-size" in inline:
        m = re.search(r"font-size\s*:\s*(\d+)", inline)
        if m:
            style["size"] = int(m.group(1))

    tag = el.name if el.name else ""
    if tag in ("b", "strong"):
        style["bold"] = True
    if tag in ("i", "em"):
        style["italic"] = True
    if tag == "u":
        style["underline"] = True

    return style


def apply_run_style(run, style):
    """应用文本样式到 run"""
    if style.get("bold"):
        run.bold = True
    if style.get("italic"):
        run.italic = True
    if style.get("underline"):
        run.underline = True
    if style.get("size"):
        run.font.size = Pt(style["size"])
    color = style.get("color")
    if color and color not in ("inherit", "initial", "unset"):
        try:
            if color.startswith("#"):
                hex_c = color[1:]
                if len(hex_c) == 3:
                    hex_c = "".join(c * 2 for c in hex_c)
                if len(hex_c) >= 6:
                    run.font.color.rgb = RGBColor(
                        int(hex_c[0:2], 16),
                        int(hex_c[2:4], 16),
                        int(hex_c[4:6], 16),
                    )
        except (ValueError, IndexError):
            pass


def clean_text(text):
    """清理文本中的异常空格和零宽字符"""
    text = text.replace("\xa0", " ")
    text = text.replace("​", "")
    text = text.replace("‌", "")
    text = text.replace("‍", "")
    text = text.replace("﻿", "")
    text = text.replace("　", " ")
    return text


# ==================== 内容处理 ====================

def process_inline(el, paragraph, doc, parent_style=None, downloaded=None,
                   img_counter=None, errors=None, cache=None):
    """递归处理内联元素，将文字和图片添加到段落"""
    if parent_style is None:
        parent_style = {}
    if downloaded is None:
        downloaded = set()
    if errors is None:
        errors = []

    for child in el.children:
        if isinstance(child, NavigableString):
            text = clean_text(str(child))
            if text.strip():
                run = paragraph.add_run(text)
                apply_run_style(run, parent_style)
        elif isinstance(child, Tag):
            tag = child.name

            # 图片
            if tag == "img":
                src = get_img_src(child)
                if src:
                    try:
                        img_data, fmt, _ = download_image(src, downloaded, cache=cache)
                        add_image_to_doc(doc, img_data, fmt, img_counter=img_counter)
                        paragraph = doc.add_paragraph()
                    except Exception as e:
                        errors.append(f"图片下载失败: {e}")
                continue

            # 换行
            if tag == "br":
                paragraph.add_run("\n")
                continue

            # 代码
            if tag in ("code", "pre"):
                code_text = child.get_text()
                run = paragraph.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 80, 80)
                continue

            # 超链接
            if tag == "a":
                href = child.get("href", "")
                link_text = child.get_text(strip=True)
                if href and link_text and not href.startswith("javascript:"):
                    add_hyperlink(paragraph, link_text, href)
                elif link_text:
                    run = paragraph.add_run(link_text)
                    apply_run_style(run, parent_style)
                continue

            # 合并样式并递归
            cur_style = dict(parent_style)
            cur_style.update(get_text_style(child))
            process_inline(child, paragraph, doc, cur_style, downloaded,
                           img_counter, errors, cache=cache)

    return paragraph


def process_element(el, doc, downloaded=None, img_counter=None, errors=None, cache=None):
    """处理一个顶层内容元素"""
    if downloaded is None:
        downloaded = set()
    if errors is None:
        errors = []

    if isinstance(el, NavigableString):
        text = clean_text(str(el)).strip()
        if text:
            doc.add_paragraph(text)
        return

    if not isinstance(el, Tag):
        return

    tag = el.name

    # 图片
    if tag == "img":
        src = get_img_src(el)
        if src:
            try:
                img_data, fmt, _ = download_image(src, downloaded, cache=cache)
                add_image_to_doc(doc, img_data, fmt, img_counter=img_counter)
            except Exception as e:
                errors.append(f"图片下载失败: {e}")
        return

    # 标题
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        text = clean_text(el.get_text(strip=True))
        if text:
            doc.add_heading(text, level=min(level, 4))
        return

    # 列表
    if tag in ("ul", "ol"):
        for item in el.find_all("li", recursive=False):
            style = "List Bullet" if tag == "ul" else "List Number"
            try:
                p = doc.add_paragraph(style=style)
            except Exception:
                p = doc.add_paragraph()
            process_inline(item, p, doc, downloaded=downloaded,
                           img_counter=img_counter, errors=errors, cache=cache)
        return

    # 引用块
    if tag == "blockquote":
        text = clean_text(el.get_text(strip=True))
        if text:
            p = doc.add_paragraph(style="Intense Quote")
            p.add_run(text)
        return

    # 表格
    if tag == "table":
        rows = el.find_all("tr")
        if rows:
            cols = max((len(r.find_all(["td", "th"])) for r in rows), default=1)
            table = doc.add_table(rows=len(rows), cols=max(cols, 1))
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for ci, cell in enumerate(cells):
                    if ci < cols:
                        table.rows[ri].cells[ci].text = clean_text(cell.get_text(strip=True))
        return

    # 分隔线
    if tag == "hr":
        doc.add_paragraph("─" * 40)
        return

    # 段落
    if tag == "p":
        imgs = el.find_all("img")
        texts = el.get_text(strip=True)
        if imgs and not texts:
            for img in imgs:
                src = get_img_src(img)
                if src:
                    try:
                        img_data, fmt, _ = download_image(src, downloaded, cache=cache)
                        add_image_to_doc(doc, img_data, fmt, img_counter=img_counter)
                    except Exception as e:
                        errors.append(f"图片下载失败: {e}")
            return
        if texts or imgs:
            p = doc.add_paragraph()
            process_inline(el, p, doc, downloaded=downloaded,
                           img_counter=img_counter, errors=errors, cache=cache)
        return

    # 容器标签 -> 递归
    if tag in ("section", "div", "article", "main", "span",
               "fieldset", "figure", "figcaption", "header", "footer"):
        for child in el.children:
            if isinstance(child, Tag):
                process_element(child, doc, downloaded, img_counter, errors, cache=cache)
            elif isinstance(child, NavigableString):
                text = clean_text(str(child)).strip()
                if text:
                    doc.add_paragraph(text)
        return

    # 其他标签 -> 段落处理
    text = clean_text(el.get_text(strip=True))
    if text:
        p = doc.add_paragraph()
        process_inline(el, p, doc, downloaded=downloaded,
                       img_counter=img_counter, errors=errors, cache=cache)


# ==================== 索引管理 ====================

def load_index(save_dir):
    """加载索引文件"""
    index_path = os.path.join(save_dir, "index.json")
    if os.path.exists(index_path):
        try:
            with open(index_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    return []


def save_index(save_dir, entries):
    """保存索引文件"""
    index_path = os.path.join(save_dir, "index.json")
    try:
        with open(index_path, "w", encoding="utf-8") as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)
    except IOError:
        pass


def record_to_index(save_dir, title, author, date, url, file_path):
    """记录文章到索引"""
    entries = load_index(save_dir)
    entries.append({
        "title": title,
        "author": author,
        "date": date,
        "url": url,
        "file": os.path.basename(file_path),
        "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    })
    save_index(save_dir, entries)


def check_duplicate(save_dir, url):
    """检查是否已保存过相同链接"""
    entries = load_index(save_dir)
    for e in entries:
        if e.get("url") == url:
            return e.get("file", "")
    return ""


# ==================== 主生成流程 ====================

def sanitize_filename(name, max_len=120):
    """清理文件名中的非法字符"""
    name = re.sub(r'[\\/:*?"<>|\r\n]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_ ')
    return name[:max_len] if name else "未命名"


def generate_docx(url, save_dir, progress_callback=None):
    """
    主流程：获取文章 -> 解析 -> 生成 Word

    Args:
        url: 公众号文章链接
        save_dir: 保存目录
        progress_callback: 进度回调 fn(msg, percent)

    Returns:
        (save_path, title) 保存的文件路径和文章标题
    """
    downloaded = set()  # 图片去重
    img_counter = [0]   # 图片计数
    errors = []         # 收集非致命错误

    def report(msg, pct):
        if progress_callback:
            progress_callback(msg, pct)

    # 0. 检查重复
    existing_file = check_duplicate(save_dir, url)
    if existing_file:
        raise ValueError(f"该文章已保存过：{existing_file}\n如需重新下载，请先删除索引中的记录。")

    # 1. 获取页面
    report("正在访问文章页面...", 5)
    soup = fetch_article(url)

    # 2. 提取信息
    report("正在解析文章信息...", 15)
    title = extract_title(soup)
    author = extract_author(soup)
    date = extract_date(soup)

    # 3. 提取正文
    report("正在提取正文内容...", 25)
    content = extract_content_area(soup)
    if not content:
        raise ValueError(
            "无法找到文章正文区域。可能原因：\n"
            "  1. 链接不是公众号文章\n"
            "  2. 文章已被删除\n"
            "  3. 需要在微信中打开才能访问"
        )

    # 4. 清理正文
    report("正在清理正文内容...", 30)
    content = clean_content_area(content)

    # 4.5 批量下载图片
    report("正在收集图片链接...", 32)
    img_urls = collect_image_urls(content)
    if img_urls:
        report(f"正在并发下载 {len(img_urls)} 张图片...", 33)
        image_cache = batch_download_images(img_urls)
        ok_count = sum(1 for v in image_cache.values() if v[0] is not None)
        report(f"图片下载完成（成功 {ok_count}/{len(img_urls)}）", 38)
    else:
        image_cache = {}

    # 5. 创建文档
    report("正在创建 Word 文档...", 35)
    doc = Document()
    setup_doc_styles(doc)

    # 标题
    doc.add_heading(title, level=0)

    # 元信息区
    meta_parts = []
    if author:
        meta_parts.append(f"作者/公众号：{author}")
    if date:
        meta_parts.append(f"发布日期：{date}")
    meta_parts.append(f"保存时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("  |  ".join(meta_parts))
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    # 原文链接行
    link_para = doc.add_paragraph()
    link_label = link_para.add_run("原文链接：")
    link_label.font.size = Pt(9)
    link_label.font.color.rgb = RGBColor(128, 128, 128)
    add_hyperlink(link_para, url, url)

    doc.add_paragraph("─" * 40)

    # 6. 处理正文
    report("正在处理正文内容...", 40)
    elements = [child for child in content.children if isinstance(child, Tag)]
    total = max(len(elements), 1)

    for i, el in enumerate(elements):
        pct = 40 + int(50 * i / total)
        if i % 5 == 0 or i == total - 1:
            report(f"正在处理正文 ({i+1}/{total})...", pct)
        try:
            process_element(el, doc, downloaded, img_counter, errors, cache=image_cache)
        except Exception as e:
            errors.append(f"段落处理异常: {e}")

    # 7. 末尾原文链接
    doc.add_paragraph("")
    doc.add_paragraph("─" * 40)
    footer = doc.add_paragraph()
    footer_label = footer.add_run("原文链接：")
    footer_label.font.size = Pt(10)
    footer_label.font.bold = True
    add_hyperlink(footer, url, url)

    # 8. 保存文件
    report("正在保存文件...", 95)

    # 规范化文件名
    parts = []
    if date:
        parts.append(date)
    if author:
        parts.append(sanitize_filename(author, 30))
    parts.append(sanitize_filename(title, 80))
    file_name = "_".join(parts) + ".docx"

    save_path = os.path.join(save_dir, file_name)

    # 文件冲突处理
    base, ext = os.path.splitext(save_path)
    counter = 1
    final_path = save_path
    while os.path.exists(final_path):
        final_path = f"{base}_{counter}{ext}"
        counter += 1

    doc.save(final_path)

    # 9. 记录索引
    record_to_index(save_dir, title, author, date, url, final_path)

    # 报告非致命错误
    if errors:
        unique_errors = list(dict.fromkeys(errors))[:10]
        err_summary = "；".join(unique_errors)
        report(f"保存完成（有 {len(errors)} 个非致命问题: {err_summary}）", 100)
    else:
        report("保存完成！", 100)

    return final_path, title


# ==================== 下载线程 ====================

class DownloadThread(QThread):
    """后台下载线程"""
    progress = pyqtSignal(str, int)  # msg, percent
    finished = pyqtSignal(str, str)  # file_path, title
    error = pyqtSignal(str)          # error message

    def __init__(self, url, save_dir):
        super().__init__()
        self.url = url
        self.save_dir = save_dir

    def run(self):
        try:
            file_path, title = generate_docx(
                self.url, self.save_dir,
                progress_callback=lambda m, p: self.progress.emit(m, p)
            )
            self.finished.emit(file_path, title)
        except Exception as e:
            self.error.emit(f"{type(e).__name__}: {e}\n{traceback.format_exc()}")


# ==================== 配置管理 ====================

CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")


def load_config():
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass
    return {}


def save_config(cfg):
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except IOError:
        pass


# ==================== GUI ====================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("公众号文章 → Word 保存工具")
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.png")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        self.setMinimumSize(780, 650)
        self.resize(850, 750)
        self.download_thread = None
        self.batch_queue = []       # 批量下载队列
        self.last_saved_file = None # 最近保存的文件
        self.init_ui()
        self.init_shortcuts()
        self.load_history()

    def init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(8)
        layout.setContentsMargins(18, 14, 18, 14)

        # 标题
        title_label = QLabel("公众号文章 → Word 保存工具")
        title_label.setFont(QFont("Microsoft YaHei", 15, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 4px 0 8px 0;")
        layout.addWidget(title_label)

        # ---- 文章信息 ----
        info_group = QGroupBox("文章信息")
        info_layout = QVBoxLayout(info_group)
        info_layout.setSpacing(6)

        # 链接输入（多行，支持批量）
        url_row = QHBoxLayout()
        url_label = QLabel("文章链接：")
        url_label.setFixedWidth(75)
        url_label.setFont(QFont("Microsoft YaHei", 10))
        url_label.setAlignment(Qt.AlignTop)
        self.url_input = QTextEdit()
        self.url_input.setPlaceholderText("粘贴公众号文章链接，每行一个\n支持批量下载，如：\nhttps://mp.weixin.qq.com/s/xxx\nhttps://mp.weixin.qq.com/s/yyy")
        self.url_input.setFont(QFont("Microsoft YaHei", 10))
        self.url_input.setMaximumHeight(90)
        url_row.addWidget(url_label)
        url_row.addWidget(self.url_input)
        info_layout.addLayout(url_row)

        # 剪贴板按钮行
        clip_row = QHBoxLayout()
        clip_row.addStretch()
        paste_btn = QPushButton("从剪贴板读取")
        paste_btn.setFixedHeight(28)
        paste_btn.setFont(QFont("Microsoft YaHei", 9))
        paste_btn.setStyleSheet("""
            QPushButton {
                background-color: #8e44ad; color: white;
                border: none; border-radius: 4px; padding: 0 14px;
            }
            QPushButton:hover { background-color: #9b59b6; }
        """)
        paste_btn.clicked.connect(self.paste_from_clipboard)
        clip_row.addWidget(paste_btn)
        info_layout.addLayout(clip_row)

        # 保存路径
        save_row = QHBoxLayout()
        save_label = QLabel("保存位置：")
        save_label.setFixedWidth(75)
        save_label.setFont(QFont("Microsoft YaHei", 10))
        self.save_path_input = QLineEdit()
        self.save_path_input.setPlaceholderText("选择保存文件夹...")
        self.save_path_input.setFont(QFont("Microsoft YaHei", 10))
        cfg = load_config()
        default_dir = cfg.get("save_dir") or os.path.dirname(os.path.abspath(__file__))
        self.save_path_input.setText(default_dir)
        browse_btn = QPushButton("浏览...")
        browse_btn.setFixedWidth(75)
        browse_btn.setFont(QFont("Microsoft YaHei", 9))
        browse_btn.clicked.connect(self.browse_folder)
        save_row.addWidget(save_label)
        save_row.addWidget(self.save_path_input)
        save_row.addWidget(browse_btn)
        info_layout.addLayout(save_row)

        layout.addWidget(info_group)

        # ---- 操作与进度 ----
        op_group = QGroupBox("操作与进度")
        op_layout = QVBoxLayout(op_group)
        op_layout.setSpacing(8)

        # 按钮行
        btn_row = QHBoxLayout()
        btn_row.addStretch()

        self.download_btn = QPushButton("  开始下载保存  ")
        self.download_btn.setFont(QFont("Microsoft YaHei", 11, QFont.Bold))
        self.download_btn.setFixedHeight(38)
        self.download_btn.setMinimumWidth(170)
        self.download_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60; color: white;
                border: none; border-radius: 5px; padding: 0 20px;
            }
            QPushButton:hover { background-color: #2ecc71; }
            QPushButton:pressed { background-color: #1e8449; }
            QPushButton:disabled { background-color: #95a5a6; }
        """)
        self.download_btn.clicked.connect(self.start_download)

        self.open_file_btn = QPushButton("打开文件")
        self.open_file_btn.setFixedHeight(38)
        self.open_file_btn.setFixedWidth(85)
        self.open_file_btn.setFont(QFont("Microsoft YaHei", 9))
        self.open_file_btn.setEnabled(False)
        self.open_file_btn.setStyleSheet("""
            QPushButton {
                background-color: #e67e22; color: white;
                border: none; border-radius: 5px;
            }
            QPushButton:hover { background-color: #f39c12; }
            QPushButton:disabled { background-color: #bdc3c7; }
        """)
        self.open_file_btn.clicked.connect(self.open_last_file)

        self.open_dir_btn = QPushButton("打开目录")
        self.open_dir_btn.setFixedHeight(38)
        self.open_dir_btn.setFixedWidth(85)
        self.open_dir_btn.setFont(QFont("Microsoft YaHei", 9))
        self.open_dir_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db; color: white;
                border: none; border-radius: 5px;
            }
            QPushButton:hover { background-color: #5dade2; }
        """)
        self.open_dir_btn.clicked.connect(self.open_save_dir)

        self.clear_btn = QPushButton("清空")
        self.clear_btn.setFixedHeight(38)
        self.clear_btn.setFixedWidth(70)
        self.clear_btn.setFont(QFont("Microsoft YaHei", 9))
        self.clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c; color: white;
                border: none; border-radius: 5px;
            }
            QPushButton:hover { background-color: #ec7063; }
        """)
        self.clear_btn.clicked.connect(self.clear_all)

        btn_row.addWidget(self.download_btn)
        btn_row.addWidget(self.open_file_btn)
        btn_row.addWidget(self.open_dir_btn)
        btn_row.addWidget(self.clear_btn)
        btn_row.addStretch()
        op_layout.addLayout(btn_row)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedHeight(22)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setFormat("就绪")
        self.progress_bar.setValue(0)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #bdc3c7; border-radius: 4px;
                text-align: center; font-size: 11px; font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #3498db; border-radius: 3px;
            }
        """)
        op_layout.addWidget(self.progress_bar)

        layout.addWidget(op_group)

        # ---- 下方区域：日志 + 历史（用 Splitter 分割）----
        splitter = QSplitter(Qt.Vertical)

        # 运行日志
        log_widget = QWidget()
        log_layout = QVBoxLayout(log_widget)
        log_layout.setContentsMargins(0, 0, 0, 0)
        log_label = QLabel("运行日志")
        log_label.setFont(QFont("Microsoft YaHei", 9, QFont.Bold))
        log_label.setStyleSheet("color: #7f8c8d;")
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setFont(QFont("Consolas", 9))
        self.log_area.setStyleSheet("""
            QTextEdit {
                background-color: #2c3e50; color: #ecf0f1;
                border: 1px solid #34495e; border-radius: 4px; padding: 6px;
            }
        """)
        log_layout.addWidget(log_label)
        log_layout.addWidget(self.log_area)
        splitter.addWidget(log_widget)

        # 历史记录
        history_widget = QWidget()
        history_layout = QVBoxLayout(history_widget)
        history_layout.setContentsMargins(0, 0, 0, 0)
        history_label = QLabel("下载历史（双击打开文件）")
        history_label.setFont(QFont("Microsoft YaHei", 9, QFont.Bold))
        history_label.setStyleSheet("color: #7f8c8d;")
        self.history_list = QListWidget()
        self.history_list.setFont(QFont("Microsoft YaHei", 9))
        self.history_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #bdc3c7; border-radius: 4px;
                background-color: #fafafa;
            }
            QListWidget::item {
                padding: 4px 6px;
                border-bottom: 1px solid #ecf0f1;
            }
            QListWidget::item:hover {
                background-color: #ebf5fb;
            }
        """)
        self.history_list.itemDoubleClicked.connect(self.open_history_file)
        history_layout.addWidget(history_label)
        history_layout.addWidget(self.history_list)
        splitter.addWidget(history_widget)

        splitter.setSizes([300, 200])
        layout.addWidget(splitter)

        self.log("就绪，请输入公众号文章链接。  快捷键：Enter 下载 | Ctrl+V 粘贴识别链接")

    def init_shortcuts(self):
        """初始化快捷键"""
        QShortcut(QKeySequence("Return"), self, self.start_download)
        QShortcut(QKeySequence("Enter"), self, self.start_download)
        QShortcut(QKeySequence("Ctrl+V"), self, self.paste_from_clipboard)

    # ---- 剪贴板 ----

    def paste_from_clipboard(self):
        """从剪贴板读取微信链接"""
        clipboard = QApplication.clipboard()
        text = clipboard.text()
        if not text:
            self.log("剪贴板为空", "warn")
            return
        # 提取所有微信链接
        urls = re.findall(r'https?://mp\.weixin\.qq\.com/s/\S+', text)
        if not urls:
            # 尝试匹配任意 http 链接
            urls = re.findall(r'https?://\S+', text)
        if urls:
            existing = self.url_input.toPlainText().strip()
            new_text = "\n".join(urls)
            if existing:
                self.url_input.setPlainText(existing + "\n" + new_text)
            else:
                self.url_input.setPlainText(new_text)
            self.log(f"从剪贴板识别到 {len(urls)} 个链接")
        else:
            self.log("剪贴板中未识别到有效链接", "warn")

    # ---- 历史记录 ----

    def load_history(self):
        """加载历史记录到列表"""
        save_dir = self.save_path_input.text().strip()
        if not save_dir or not os.path.isdir(save_dir):
            return
        entries = load_index(save_dir)
        self.history_list.clear()
        for entry in reversed(entries):  # 最新的在前
            title = entry.get("title", "未知")
            author = entry.get("author", "")
            date = entry.get("date", "")
            file_name = entry.get("file", "")
            display = f"[{date}] {author} — {title}" if author else f"[{date}] {title}"
            item = QListWidgetItem(display)
            item.setData(Qt.UserRole, os.path.join(save_dir, file_name))
            self.history_list.addItem(item)

    def open_history_file(self, item):
        """双击打开历史文件"""
        file_path = item.data(Qt.UserRole)
        if file_path and os.path.exists(file_path):
            try:
                os.startfile(file_path)
            except AttributeError:
                subprocess.Popen(["xdg-open", file_path])
        else:
            QMessageBox.warning(self, "提示", f"文件不存在：\n{file_path}")

    # ---- 日志 ----

    def log(self, msg, level="info"):
        timestamp = datetime.now().strftime("%H:%M:%S")
        level_tag = {"info": "INFO", "warn": "WARN", "error": "ERROR"}.get(level, "INFO")
        color = {"info": "#ecf0f1", "warn": "#f39c12", "error": "#e74c3c"}.get(level, "#ecf0f1")
        html = (
            f'<span style="color:#95a5a6">[{timestamp}]</span> '
            f'<span style="color:{color}">[{level_tag}] {msg}</span>'
        )
        self.log_area.append(html)
        sb = self.log_area.verticalScrollBar()
        sb.setValue(sb.maximum())

    # ---- 路径操作 ----

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择保存文件夹")
        if folder:
            self.save_path_input.setText(folder)
            save_config({"save_dir": folder})
            self.load_history()

    def open_save_dir(self):
        path = self.save_path_input.text().strip()
        if path and os.path.isdir(path):
            try:
                os.startfile(path)
            except AttributeError:
                subprocess.Popen(["xdg-open", path])
        else:
            QMessageBox.warning(self, "提示", "保存路径不存在，请先选择有效路径。")

    def open_last_file(self):
        """打开最近保存的文件"""
        if self.last_saved_file and os.path.exists(self.last_saved_file):
            try:
                os.startfile(self.last_saved_file)
            except AttributeError:
                subprocess.Popen(["xdg-open", self.last_saved_file])

    def clear_all(self):
        self.url_input.clear()
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat("就绪")
        self.log_area.clear()
        self.batch_queue.clear()
        self.open_file_btn.setEnabled(False)
        self.log("已清空，请输入新的链接。")

    # ---- 下载逻辑 ----

    def get_urls_from_input(self):
        """从输入框提取所有链接"""
        text = self.url_input.toPlainText().strip()
        if not text:
            return []
        urls = []
        for line in text.splitlines():
            line = line.strip()
            if line and line.startswith("http"):
                urls.append(line)
        return urls

    def validate_inputs(self):
        urls = self.get_urls_from_input()
        if not urls:
            QMessageBox.warning(self, "提示", "请输入公众号文章链接！")
            self.url_input.setFocus()
            return False, []

        save_dir = self.save_path_input.text().strip()
        if not save_dir:
            QMessageBox.warning(self, "提示", "请选择保存文件夹！")
            return False, []

        if not os.path.isdir(save_dir):
            QMessageBox.warning(self, "提示", f"保存路径不存在：\n{save_dir}")
            return False, []

        if not os.access(save_dir, os.W_OK):
            QMessageBox.warning(self, "提示", f"保存路径不可写，请选择其他目录：\n{save_dir}")
            return False, []

        # 保存路径记忆
        save_config({"save_dir": save_dir})

        return True, urls

    def set_busy(self, busy):
        self.download_btn.setEnabled(not busy)
        self.clear_btn.setEnabled(not busy)

    def start_download(self):
        # 如果正在下载中，忽略
        if self.download_thread and self.download_thread.isRunning():
            return

        valid, urls = self.validate_inputs()
        if not valid:
            return

        self.batch_queue = list(urls)
        self.set_busy(True)
        self.open_file_btn.setEnabled(False)
        self.log(f"共 {len(self.batch_queue)} 篇文章待下载")
        self._start_next_download()

    def _start_next_download(self):
        """启动队列中的下一个下载"""
        if not self.batch_queue:
            self.set_busy(False)
            self.progress_bar.setValue(100)
            self.progress_bar.setFormat("全部完成")
            self.log("所有文章下载完成！")
            self.load_history()
            return

        url = self.batch_queue.pop(0)
        remaining = len(self.batch_queue)
        save_dir = self.save_path_input.text().strip()

        self.progress_bar.setValue(0)
        self.progress_bar.setFormat(f"[剩余 {remaining + 1} 篇] 准备中...")
        self.log(f"开始处理: {url}")

        self.download_thread = DownloadThread(url, save_dir)
        self.download_thread.progress.connect(self.on_progress)
        self.download_thread.finished.connect(self.on_finished)
        self.download_thread.error.connect(self.on_error)
        self.download_thread.start()

    def on_progress(self, msg, pct):
        self.progress_bar.setValue(pct)
        remaining = len(self.batch_queue)
        prefix = f"[剩余 {remaining + 1} 篇] " if remaining > 0 else ""
        self.progress_bar.setFormat(f"{prefix}{pct}% — {msg}")
        level = "warn" if "失败" in msg or "警告" in msg else "info"
        self.log(msg, level)

    def on_finished(self, file_path, title):
        self.last_saved_file = file_path
        self.open_file_btn.setEnabled(True)
        self.log(f"文章「{title}」已保存: {file_path}")
        # 继续下一篇
        self._start_next_download()

    def on_error(self, err_msg):
        self.log(err_msg.split("\n")[0], "error")
        # 继续下一篇
        self._start_next_download()


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    # 设置应用图标（任务栏）
    icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.png")
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))

    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(245, 246, 250))
    palette.setColor(QPalette.WindowText, QColor(44, 62, 80))
    palette.setColor(QPalette.Base, QColor(255, 255, 255))
    palette.setColor(QPalette.AlternateBase, QColor(245, 246, 250))
    palette.setColor(QPalette.ToolTipBase, QColor(44, 62, 80))
    palette.setColor(QPalette.ToolTipText, QColor(255, 255, 255))
    palette.setColor(QPalette.Text, QColor(44, 62, 80))
    palette.setColor(QPalette.Button, QColor(245, 246, 250))
    palette.setColor(QPalette.ButtonText, QColor(44, 62, 80))
    palette.setColor(QPalette.Highlight, QColor(52, 152, 219))
    palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
    app.setPalette(palette)

    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
